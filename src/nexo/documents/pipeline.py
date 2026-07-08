"""File processing pipeline: extract text -> summarize.

Wired into the WeCom file route via `nexo.app.handle_file`. The transport
layer (wecom.py) owns download + AES decryption via the SDK; this module only
sees the already-decrypted file bytes. The pipeline is an `AsyncIterator[str]`
so each stage yields a progress chunk that the transport streams back to the
WeCom bubble:

    "正在解析…" -> "正在生成摘要…" -> <summary markdown>

The original upload is persisted to Huawei Cloud OBS (`nexo.storage.obs`);
text is extracted in-memory (no local disk write). The generated summary is
NOT persisted anywhere — it only appears as the WeCom reply.

DeepSeek (the configured OpenAI-compatible model) has no file-upload endpoint,
so text is always extracted client-side before being handed to `ingest_agent`.
"""

from __future__ import annotations

import io
from collections.abc import AsyncIterator
from pathlib import Path

from nexo.agents.ingest import IngestResult, run as ingest_agent_run
from nexo.prompts import msg
from nexo.storage.obs import upload_upload


async def process_file(file_data: bytes, filename: str) -> AsyncIterator[str]:
    """Persist the original to OBS, extract text in memory, summarize, reply.

    `file_data` is the already-downloaded-and-decrypted file content (the
    transport layer handles download + AES decryption). The original is uploaded
    to OBS; text is extracted from an in-memory buffer. Yields progress messages
    and finally the summary markdown. The summary is not persisted — it only
    surfaces as the WeCom reply.
    """
    if not file_data:
        raise ValueError("文件内容为空，无法处理")

    # 1. Persist the original to OBS (durable store; nothing written locally).
    try:
        await upload_upload(filename, file_data)
    except Exception as exc:  # noqa: BLE001 — surface a readable message, don't hang the bubble
        yield msg("file_save_failed", error=exc)
        return

    # 2. Extract text in memory (no local tempfile).
    name = Path(filename).name or filename or "unnamed"
    yield msg("file_parsing", name=name)
    text = _extract_text(name, file_data)
    if not text.strip():
        yield msg("file_extract_empty")
        return

    # 3. Summarize via the existing ingest agent.
    yield msg("file_summarizing")
    result = await ingest_agent_run(text)

    # 4. Reply with the summary (no persistence — it only shows in the bubble).
    yield _format_markdown(result)


def _extract_text(name: str, data: bytes) -> str:
    """Extract plain text from in-memory bytes, dispatching by extension."""
    suffix = Path(name).suffix.lower()
    stream = io.BytesIO(data)
    if suffix == ".pdf":
        return _extract_pdf(stream)
    if suffix == ".docx":
        return _extract_docx(stream)
    if suffix == ".doc":
        raise NotImplementedError("旧版 .doc 暂不支持，请另存为 .docx 后重试")
    if suffix in (".txt", ".md", ".markdown", ".csv", ".log"):
        return _read_text(data)
    raise ValueError(f"不支持的文件类型: {suffix or '(无扩展名)'}")


def _extract_pdf(stream: io.BytesIO) -> str:
    from pypdf import PdfReader

    reader = PdfReader(stream)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(stream: io.BytesIO) -> str:
    import docx  # python-docx

    document = docx.Document(stream)
    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_text(data: bytes) -> str:
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def _format_markdown(result: IngestResult) -> str:
    """Render an IngestResult as a self-contained summary markdown document,
    following the template: 📌 核心摘要 / 🎯 核心要点 / 🔑 关键词, plus the
    optional 🎬 背景 and 🛠️ 下步行动 sections (only when the agent filled them).
    """
    out: list[str] = [
        "## 📌 一句话核心摘要",
        "",
        result.core_summary,
        "",
        "## 🎯 核心要点提炼",
        "",
    ]
    for kp in result.key_points:
        out.append(f"- **{kp.headline}**：{kp.detail}")

    out += ["", "## 🔑 关键词/标签", ""]
    out.append(
        " ".join(f"#{t}" for t in result.tags) if result.tags else msg("no_keywords")
    )

    if result.background:
        out += ["", "## 🎬 背景/上下文", "", result.background]

    if result.action_items:
        out += ["", "## 🛠️ 下步行动/待办事项", ""]
        out += [f"- {a}" for a in result.action_items]

    out.append("")
    return "\n".join(out)
